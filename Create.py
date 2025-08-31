#!/usr/bin/env python3
"""Überstiegstest Generator v2.0 Vollständig überarbeitete Version mit allen
Anforderungen."""

import argparse
import ast
import json
import math
import operator
import random
import sys
from dataclasses import dataclass
from enum import Enum
from fractions import Fraction

# Für Word-Export
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    # Dummy-Klassen für den Fall, dass python-docx nicht installiert ist
    Document = None
    Pt = None

# Für LaTeX-Export
try:
    import subprocess

    HAS_LATEX = (
        subprocess.run(["pdflatex", "--version"], capture_output=True).returncode == 0
    )
except Exception:
    HAS_LATEX = False


def fmt(x: float, nd: int = 2, thousand: bool = False) -> str:
    s = f"{x:,.{nd}f}" if thousand else f"{x:.{nd}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_int_or_dec(x: float, nd_if_dec: int = 2) -> str:
    return f"{int(x)}" if float(x).is_integer() else fmt(x, nd_if_dec)


class Schwierigkeit(Enum):
    """Schwierigkeitsgrade für den Test."""

    EINFACH = 1
    MITTEL = 2
    SCHWER = 3


@dataclass
class AustrianData:
    """Realistische österreichische Daten für Textaufgaben."""

    berufe = {
        "elektriker": {"gehalt_min": 2000, "gehalt_max": 2600, "stunden": (38, 42)},
        "mechaniker": {"gehalt_min": 1900, "gehalt_max": 2400, "stunden": (38, 40)},
        "schweisser": {"gehalt_min": 2100, "gehalt_max": 2500, "stunden": (38, 40)},
        "kfz_techniker": {"gehalt_min": 1800, "gehalt_max": 2300, "stunden": (38, 42)},
        "maurer": {"gehalt_min": 2200, "gehalt_max": 2700, "stunden": (38, 40)},
        "installateur": {"gehalt_min": 2000, "gehalt_max": 2500, "stunden": (38, 40)},
        "cnc_fraesen": {"gehalt_min": 2300, "gehalt_max": 2800, "stunden": (38, 42)},
    }

    preise = {
        "stahl_kg": 0.85,
        "aluminium_kg": 2.20,
        "kupfer_kg": 8.50,
        "strom_kwh": 0.22,
        "benzin_l": 1.45,
        "diesel_l": 1.42,
        "holz_m3": 120,
        "beton_m3": 95,
        "wasser_m3": 3.80,
        "gas_kwh": 0.08,
    }

    firmen = [
        "voestalpine",
        "Siemens Österreich",
        "Magna Steyr",
        "AVL List",
        "Andritz AG",
        "Palfinger",
        "Rosenbauer",
        "KTM",
        "Doppelmayr",
    ]

    materialien = {
        "stahl": {"dichte": 7.85, "einheit": "kg/dm³"},
        "aluminium": {"dichte": 2.70, "einheit": "kg/dm³"},
        "kupfer": {"dichte": 8.96, "einheit": "kg/dm³"},
        "holz_fichte": {"dichte": 0.47, "einheit": "kg/dm³"},
        "beton": {"dichte": 2.40, "einheit": "kg/dm³"},
    }


class MathSolver:
    """Robuster mathematischer Solver mit Bruchrechnung."""

    _ops = {
        ast.Add: operator.add,
        ast.Sub: operator.sub,
        ast.Mult: operator.mul,
        ast.Div: operator.truediv,
    }

    @staticmethod
    def evaluate_expression(expr: str) -> float | None:
        """Sichere Auswertung mathematischer Ausdrücke."""

        def _eval(node):
            if isinstance(node, (ast.Num, ast.Constant)):
                v = getattr(node, "n", getattr(node, "value", None))
                if not isinstance(v, (int, float)):
                    raise ValueError("Ungültiger Wert")
                return v
            if isinstance(node, ast.UnaryOp) and isinstance(node.op, ast.USub):
                return -_eval(node.operand)
            if isinstance(node, ast.BinOp) and type(node.op) in MathSolver._ops:
                return MathSolver._ops[type(node.op)](
                    _eval(node.left), _eval(node.right)
                )
            raise ValueError("Ungültiger Ausdruck")

        try:
            expr = (
                expr.replace("·", "*")
                .replace("×", "*")
                .replace(":", "/")
                .replace("÷", "/")
                .replace(",", ".")
            )
            tree = ast.parse(expr, mode="eval")
            return round(float(_eval(tree.body)), 4)
        except Exception:
            return None

    @staticmethod
    def solve_fraction(numerator: int, denominator: int) -> Fraction:
        """Erstellt und kürzt einen Bruch."""
        return Fraction(numerator, denominator)

    @staticmethod
    def add_fractions(f1: Fraction, f2: Fraction) -> tuple[Fraction, int]:
        """Addiert zwei Brüche und gibt gemeinsamen Nenner zurück."""
        result = f1 + f2
        common_denominator = (
            f1.denominator * f2.denominator // math.gcd(f1.denominator, f2.denominator)
        )
        return result, common_denominator

    @staticmethod
    def solve_linear_equation(a: float, b: float, c: float, d: float) -> float:
        """Löst ax + b = cx + d."""
        if a == c:
            return None if b != d else "unendlich viele Lösungen"
        return (d - b) / (a - c)

    @staticmethod
    def round_to_place(value: float, place: str) -> float:
        """Rundet auf die angegebene Stelle."""
        places_map = {
            "t": 3,  # Tausendstel
            "h": 2,  # Hundertstel
            "z": 1,  # Zehntel
            "E": 0,  # Einer
            "Z": -1,  # Zehner
            "H": -2,  # Hunderter
            "T": -3,  # Tausender
            "ZT": -4,  # Zehntausender
            "HT": -5,  # Hunderttausender
            "M": -6,  # Million
        }

        if place not in places_map:
            return value

        decimals = places_map[place]

        if decimals >= 0:
            return round(value, decimals)
        else:
            # Runden auf Zehner, Hunderter, etc.
            factor = 10 ** abs(decimals)
            return round(value / factor) * factor


class QualityControl:
    """Qualitätskontrolle für generierte Aufgaben."""

    def __init__(self):
        self.used_templates = []
        self.used_numbers = []
        self.similarity_threshold = 0.7

    def check_similarity(self, numbers: list[int]) -> bool:
        """Prüft ob Zahlen zu ähnlich zu vorherigen sind."""
        if not self.used_numbers:
            return True

        for prev_numbers in self.used_numbers[-5:]:  # Check last 5
            if (
                self._calculate_similarity(numbers, prev_numbers)
                > self.similarity_threshold
            ):
                return False
        return True

    def _calculate_similarity(self, list1: list[int], list2: list[int]) -> float:
        """Berechnet Ähnlichkeit zweier Zahlenlisten."""
        if len(list1) != len(list2):
            return 0

        matches = sum(
            1
            for a, b in zip(sorted(list1), sorted(list2))
            if abs(a - b) < max(a, b) * 0.1
        )
        return matches / len(list1)

    def register_numbers(self, numbers: list[int]):
        """Registriert verwendete Zahlen."""
        self.used_numbers.append(numbers)

    def check_template(self, template_id: str) -> bool:
        """Prüft ob Template kürzlich verwendet wurde."""
        if template_id in self.used_templates[-3:]:
            return False
        return True

    def register_template(self, template_id: str):
        """Registriert verwendetes Template."""
        self.used_templates.append(template_id)

    def validate_range(
        self, value: float, min_val: float, max_val: float, name: str
    ) -> bool:
        """Validiert ob Wert in realistischem Bereich liegt."""
        if not min_val <= value <= max_val:
            print(
                f"Warnung: {name} = {value} außerhalb des Bereichs [{min_val}, {max_val}]"
            )
            return False
        return True


class GeometryCalculator:
    """Berechnet geometrische Aufgaben."""

    @staticmethod
    def rectangle_area(length: float, width: float) -> float:
        """Fläche eines Rechtecks."""
        return length * width

    @staticmethod
    def rectangle_perimeter(length: float, width: float) -> float:
        """Umfang eines Rechtecks."""
        return 2 * (length + width)

    @staticmethod
    def triangle_area(base: float, height: float) -> float:
        """Fläche eines Dreiecks."""
        return base * height / 2

    @staticmethod
    def circle_area(radius: float) -> float:
        """Fläche eines Kreises."""
        return math.pi * radius**2

    @staticmethod
    def circle_perimeter(radius: float) -> float:
        """Umfang eines Kreises."""
        return 2 * math.pi * radius

    @staticmethod
    def l_shape_area(l1: float, w1: float, l2: float, w2: float) -> float:
        """Fläche einer L-Form."""
        return l1 * w1 + l2 * w2

    @staticmethod
    def l_shape_perimeter(l1: float, w1: float, l2: float, w2: float) -> float:
        """Umfang einer L-Form."""
        pts = [
            (0, 0),
            (l1, 0),
            (l1, w1),
            (l1 - l2, w1),
            (l1 - l2, w1 - w2),
            (0, w1 - w2),
        ]
        return sum(
            math.hypot(pts[(i + 1) % 6][0] - pts[i][0], pts[(i + 1) % 6][1] - pts[i][1])
            for i in range(6)
        )

    @staticmethod
    def cuboid_volume(length: float, width: float, height: float) -> float:
        """Volumen eines Quaders."""
        return length * width * height

    @staticmethod
    def cuboid_surface(length: float, width: float, height: float) -> float:
        """Oberfläche eines Quaders."""
        return 2 * (length * width + length * height + width * height)


class UnitConverter:
    """Konvertiert zwischen verschiedenen Einheiten."""

    # Umrechnungsfaktoren zur Basiseinheit
    conversions = {
        # Länge (Basis: m)
        "mm": 0.001,
        "cm": 0.01,
        "dm": 0.1,
        "m": 1,
        "km": 1000,
        # Fläche (Basis: m²)
        "mm²": 0.000001,
        "cm²": 0.0001,
        "dm²": 0.01,
        "m²": 1,
        "km²": 1000000,
        "ha": 10000,
        "a": 100,
        # Volumen (Basis: m³)
        "mm³": 0.000000001,
        "cm³": 0.000001,
        "dm³": 0.001,
        "m³": 1,
        "ml": 0.000001,
        "cl": 0.00001,
        "dl": 0.0001,
        "l": 0.001,
        "hl": 0.1,
        # Masse (Basis: kg)
        "mg": 0.000001,
        "g": 0.001,
        "kg": 1,
        "t": 1000,
        # Zeit (Basis: s)
        "s": 1,
        "min": 60,
        "h": 3600,
        "d": 86400,
    }

    @classmethod
    def convert(cls, value: float, from_unit: str, to_unit: str) -> float | None:
        """Konvertiert zwischen Einheiten."""
        if from_unit not in cls.conversions or to_unit not in cls.conversions:
            return None

        # Konvertiere zur Basiseinheit und dann zur Zieleinheit
        base_value = value * cls.conversions[from_unit]
        result = base_value / cls.conversions[to_unit]
        return result


class AufgabenGenerator:
    """Generiert verschiedene Aufgabentypen."""

    def __init__(self, schwierigkeit: Schwierigkeit = Schwierigkeit.MITTEL):
        self.schwierigkeit = schwierigkeit
        self.quality_control = QualityControl()
        self.austrian_data = AustrianData()
        self.math_solver = MathSolver()
        self.geometry = GeometryCalculator()
        self.converter = UnitConverter()

    def _register_task(self, template_id: str, numbers: list[int]) -> bool:
        if not self.quality_control.check_similarity(numbers):
            return False
        self.quality_control.register_template(template_id)
        self.quality_control.register_numbers(numbers)
        return True

    def generate_grundrechnung(self, punkte: int) -> tuple[str, str, str]:
        """Generiert Grundrechenaufgabe."""
        if punkte <= 2:  # Leicht
            return self._grundrechnung_leicht()
        elif punkte <= 3:  # Mittel
            return self._grundrechnung_mittel()
        else:  # Schwer
            return self._grundrechnung_schwer()

    def _grundrechnung_leicht(self) -> tuple[str, str, str]:
        """Leichte Grundrechenaufgabe."""
        templates = [
            lambda: self._template_addition(),
            lambda: self._template_subtraktion(),
            lambda: self._template_multiplikation(),
            lambda: self._template_division(),
        ]
        return random.choice(templates)()

    def _template_addition(self) -> tuple[str, str, str]:
        """Addition Template."""
        while True:
            a = (
                random.randint(50, 500)
                if self.schwierigkeit == Schwierigkeit.EINFACH
                else random.randint(100, 999)
            )
            b = (
                random.randint(20, 200)
                if self.schwierigkeit == Schwierigkeit.EINFACH
                else random.randint(50, 500)
            )
            c = (
                random.randint(10, 100)
                if self.schwierigkeit == Schwierigkeit.EINFACH
                else random.randint(20, 300)
            )
            if self._register_task("grundrechnung/addition", [a, b, c]):
                break

        aufgabe = f"{a} + {b} - {c}"
        result = a + b - c
        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(a + b)}\nSchritt 2: {fmt(a + b)} - {fmt(c)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _template_subtraktion(self) -> tuple[str, str, str]:
        """Subtraktion Template."""
        while True:
            a = random.randint(500, 1000)
            b = random.randint(100, 400)
            c = random.randint(50, 200)
            if self._register_task("grundrechnung/subtraktion", [a, b, c]):
                break

        aufgabe = f"{a} - {b} - {c}"
        result = a - b - c
        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(a)} - {fmt(b)} = {fmt(a - b)}\n"
            f"Schritt 2: {fmt(a - b)} - {fmt(c)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _template_multiplikation(self) -> tuple[str, str, str]:
        """Multiplikation Template."""
        while True:
            a = random.randint(12, 25)
            b = random.randint(3, 12)
            if self._register_task("grundrechnung/multiplikation", [a, b]):
                break

        aufgabe = f"{a} · {b}"
        result = a * b
        loesung = fmt(result)
        erklaerung = f"{fmt(a)} · {fmt(b)} = {loesung}"

        return aufgabe, loesung, erklaerung

    def _template_division(self) -> tuple[str, str, str]:
        """Division Template."""
        while True:
            b = random.randint(5, 15)
            result = random.randint(10, 50)
            a = b * result
            if self._register_task("grundrechnung/division", [a, b]):
                break

        aufgabe = f"{a} : {b}"
        loesung = fmt(result)
        erklaerung = f"{fmt(a)} : {fmt(b)} = {loesung}"

        return aufgabe, loesung, erklaerung

    def _grundrechnung_mittel(self) -> tuple[str, str, str]:
        """Mittlere Grundrechenaufgabe mit Klammern."""
        templates = [
            lambda: self._template_klammer_plus(),
            lambda: self._template_klammer_minus(),
            lambda: self._template_klammer_mal(),
        ]
        return random.choice(templates)()

    def _template_klammer_plus(self) -> tuple[str, str, str]:
        """Klammer mit Addition."""
        while True:
            a = random.randint(80, 150)
            b = random.randint(10, 30)
            c = random.randint(3, 8)
            d = random.randint(5, 15)
            if self._register_task("grundrechnung/klammer_plus", [a, b, c, d]):
                break

        aufgabe = f"{a} - ({b} + {c} · {d})"
        result = a - (b + c * d)
        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(c)} · {fmt(d)} = {fmt(c * d)}\n"
            f"Schritt 2: {fmt(b)} + {fmt(c * d)} = {fmt(b + c * d)}\n"
            f"Schritt 3: {fmt(a)} - {fmt(b + c * d)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _template_klammer_minus(self) -> tuple[str, str, str]:
        """Klammer mit Subtraktion."""
        while True:
            a = random.randint(20, 40)
            b = random.randint(10, 25)
            c = random.randint(2, 6)
            d = random.randint(10, 30)
            if self._register_task("grundrechnung/klammer_minus", [a, b, c, d]):
                break

        aufgabe = f"({a} + {b}) · {c} - {d}"
        result = (a + b) * c - d
        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(a)} + {fmt(b)} = {fmt(a + b)}\n"
            f"Schritt 2: {fmt(a + b)} · {fmt(c)} = {fmt((a + b) * c)}\n"
            f"Schritt 3: {fmt((a + b) * c)} - {fmt(d)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _template_klammer_mal(self) -> tuple[str, str, str]:
        """Klammer mit Multiplikation."""
        while True:
            a = random.randint(100, 200)
            b = random.randint(5, 15)
            c = random.randint(3, 8)
            d = random.randint(20, 50)
            if self._register_task("grundrechnung/klammer_mal", [a, b, c, d]):
                break

        aufgabe = f"{a} + {b} · ({c} + {d})"
        result = a + b * (c + d)
        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(c)} + {fmt(d)} = {fmt(c + d)}\n"
            f"Schritt 2: {fmt(b)} · {fmt(c + d)} = {fmt(b * (c + d))}\n"
            f"Schritt 3: {fmt(a)} + {fmt(b * (c + d))} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _grundrechnung_schwer(self) -> tuple[str, str, str]:
        """Schwere Grundrechenaufgabe mit verschachtelten Klammern."""
        templates = [
            lambda: self._template_verschachtelt1(),
            lambda: self._template_verschachtelt2(),
            lambda: self._template_negativ(),
        ]
        return random.choice(templates)()

    def _template_verschachtelt1(self) -> tuple[str, str, str]:
        """Verschachtelte Klammern Typ 1."""
        while True:
            a = random.randint(100, 200)
            b = random.randint(8, 20)
            c = random.randint(2, 6)
            d = random.randint(10, 25)
            e = random.randint(3, 8)
            f = random.randint(10, 40)
            if self._register_task("grundrechnung/verschachtelt1", [a, b, c, d, e, f]):
                break

        aufgabe = f"{a} - [{b} + {c} · ({d} - {e})] + {f}"
        inner = d - e
        mult = c * inner
        bracket = b + mult
        result = a - bracket + f

        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(d)} - {fmt(e)} = {fmt(inner)}\n"
            f"Schritt 2: {fmt(c)} · {fmt(inner)} = {fmt(mult)}\n"
            f"Schritt 3: {fmt(b)} + {fmt(mult)} = {fmt(bracket)}\n"
            f"Schritt 4: {fmt(a)} - {fmt(bracket)} = {fmt(a - bracket)}\n"
            f"Schritt 5: {fmt(a - bracket)} + {fmt(f)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _template_verschachtelt2(self) -> tuple[str, str, str]:
        """Verschachtelte Klammern Typ 2."""
        while True:
            a = random.randint(150, 250)
            b = random.randint(20, 40)
            c = random.randint(3, 7)
            d = random.randint(5, 15)
            e = random.randint(10, 30)
            f = random.randint(2, 5)
            if self._register_task("grundrechnung/verschachtelt2", [a, b, c, d, e, f]):
                break

        aufgabe = f"[{a} - ({b} · {c})] + ({d} + {e}) : {f}"
        mult = b * c
        first_bracket = a - mult
        second_bracket = d + e
        div = second_bracket / f
        result = first_bracket + div

        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: {fmt(b)} · {fmt(c)} = {fmt(mult)}\n"
            f"Schritt 2: {fmt(a)} - {fmt(mult)} = {fmt(first_bracket)}\n"
            f"Schritt 3: {fmt(d)} + {fmt(e)} = {fmt(second_bracket)}\n"
            f"Schritt 4: {fmt(second_bracket)} : {fmt(f)} = {fmt(div)}\n"
            f"Schritt 5: {fmt(first_bracket)} + {fmt(div)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def _template_negativ(self) -> tuple[str, str, str]:
        """Mit negativen Zahlen."""
        while True:
            a = random.randint(20, 50)
            b = random.randint(2, 8)
            c = random.randint(3, 9)
            d = random.randint(15, 40)
            e = random.randint(5, 20)
            f = random.randint(2, 6)
            if self._register_task("grundrechnung/negativ", [a, b, c, d, e, f]):
                break

        aufgabe = f"-{a} + (-{b}) · {c} - ({d} + {e}) : {f}"
        neg_mult = -b * c
        sum_de = d + e
        div = sum_de / f
        result = -a + neg_mult - div

        loesung = fmt(result)
        erklaerung = (
            f"Schritt 1: (-{fmt(b)}) · {fmt(c)} = {fmt(neg_mult)}\n"
            f"Schritt 2: {fmt(d)} + {fmt(e)} = {fmt(sum_de)}\n"
            f"Schritt 3: {fmt(sum_de)} : {fmt(f)} = {fmt(div)}\n"
            f"Schritt 4: -{fmt(a)} + {fmt(neg_mult)} = {fmt(-a + neg_mult)}\n"
            f"Schritt 5: {fmt(-a + neg_mult)} - {fmt(div)} = {loesung}"
        )

        return aufgabe, loesung, erklaerung

    def generate_bruchaufgabe(self, punkte: int) -> tuple[str, str, str]:
        """Generiert Bruchaufgabe mit Nenner < 100."""
        if punkte <= 4:
            niveau = 1 if self.schwierigkeit == Schwierigkeit.EINFACH else 2
        else:
            niveau = 3

        # Generiere Brüche mit kleinen Nennern
        max_denominator = 12 if niveau == 1 else 20 if niveau == 2 else 30

        # Stelle sicher, dass Nenner unterschiedlich sind (kein gemeinsamer Nenner von Anfang an)
        d1 = random.randint(2, max_denominator)
        d2 = random.randint(2, max_denominator)
        while d2 == d1 or (d1 % d2 == 0) or (d2 % d1 == 0):
            d2 = random.randint(2, max_denominator)

        # Prüfe ob gemeinsamer Nenner < 100
        lcm = (d1 * d2) // math.gcd(d1, d2)
        if lcm >= 100:
            # Versuche kleinere Nenner
            d1 = random.randint(2, 10)
            d2 = random.randint(2, 12)
            while d2 == d1:
                d2 = random.randint(2, 12)
            lcm = (d1 * d2) // math.gcd(d1, d2)

        n1 = random.randint(1, d1 - 1)
        n2 = random.randint(1, d2 - 1)

        if niveau == 1:
            # Einfache Addition/Subtraktion
            operation = random.choice(["+", "-"])
            f1 = Fraction(n1, d1)
            f2 = Fraction(n2, d2)

            if operation == "+":
                result = f1 + f2
                aufgabe = f"{n1}/{d1} + {n2}/{d2}"
                # Erweiterte Brüche für gemeinsamen Nenner
                n1_erweitert = n1 * (lcm // d1)
                n2_erweitert = n2 * (lcm // d2)
                erklaerung = f"Gemeinsamer Nenner: {lcm} → {n1_erweitert}/{lcm} + {n2_erweitert}/{lcm} = {n1_erweitert + n2_erweitert}/{lcm}"
            else:
                # Stelle sicher, dass Ergebnis positiv ist
                if f1 < f2:
                    f1, f2 = f2, f1
                    n1, d1, n2, d2 = n2, d2, n1, d1
                result = f1 - f2
                aufgabe = f"{n1}/{d1} - {n2}/{d2}"
                n1_erweitert = n1 * (lcm // d1)
                n2_erweitert = n2 * (lcm // d2)
                erklaerung = f"Gemeinsamer Nenner: {lcm} → {n1_erweitert}/{lcm} - {n2_erweitert}/{lcm} = {n1_erweitert - n2_erweitert}/{lcm}"

            # Formatiere Lösung
            if result.denominator == 1:
                loesung = str(result.numerator)
            else:
                loesung = f"{result.numerator}/{result.denominator}"

        elif niveau == 2:
            # Mit gemischten Zahlen
            whole = random.randint(1, 3)
            f1 = Fraction(whole * d1 + n1, d1)
            f2 = Fraction(n2, d2)

            result = f1 + f2
            aufgabe = f"{whole} {n1}/{d1} + {n2}/{d2}"

            # Formatiere Lösung als gemischte Zahl
            if result.numerator >= result.denominator:
                whole_part = result.numerator // result.denominator
                rest = result.numerator % result.denominator
                if rest == 0:
                    loesung = str(whole_part)
                else:
                    loesung = f"{whole_part} {rest}/{result.denominator}"
            else:
                loesung = f"{result.numerator}/{result.denominator}"

            erklaerung = f"Umwandlung: {whole} {n1}/{d1} = {f1.numerator}/{f1.denominator}, dann Addition mit gemeinsamem Nenner {lcm}"

        else:
            # Schwer mit Multiplikation/Division
            operation = random.choice(["·", ":"])
            f1 = Fraction(n1, d1)
            f2 = Fraction(n2, d2)

            if operation == "·":
                result = f1 * f2
                aufgabe = f"({n1}/{d1}) · ({n2}/{d2})"
                erklaerung = f"Zähler mal Zähler, Nenner mal Nenner: ({n1}·{n2})/({d1}·{d2}) = {n1 * n2}/{d1 * d2}"
                # Kürzen falls möglich
                if result.numerator != n1 * n2 or result.denominator != d1 * d2:
                    erklaerung += (
                        f" = {result.numerator}/{result.denominator} (gekürzt)"
                    )
            else:
                result = f1 / f2
                aufgabe = f"({n1}/{d1}) : ({n2}/{d2})"
                erklaerung = f"Mit Kehrwert multiplizieren: ({n1}/{d1}) · ({d2}/{n2}) = {n1 * d2}/{d1 * n2}"
                # Kürzen falls möglich
                if result.numerator != n1 * d2 or result.denominator != d1 * n2:
                    erklaerung += (
                        f" = {result.numerator}/{result.denominator} (gekürzt)"
                    )

            if result.denominator == 1:
                loesung = str(result.numerator)
            else:
                loesung = f"{result.numerator}/{result.denominator}"

        return aufgabe, loesung, erklaerung

    def generate_gleichung(self, schwer: bool = False) -> tuple[str, str, str]:
        """Generiert Gleichung."""
        if not schwer:
            # Mittlere Gleichung: 3 Klammern, keine Brüche
            a1, b1, c1 = (
                random.randint(2, 5),
                random.randint(1, 4),
                random.randint(2, 10),
            )
            a2, b2 = random.randint(2, 4), random.randint(1, 8)
            a3, b3 = random.randint(1, 3), random.randint(2, 8)

            # Eine Klammer mit -, eine mit +, eine mit ·
            aufgabe = f"{a1}({b1}x + {c1}) - {a2}(x - {b2}) + {a3}·(x + {b3}) = 0"

            # Auflösen
            # a1*b1*x + a1*c1 - a2*x + a2*b2 + a3*x + a3*b3 = 0
            coeff_x = a1 * b1 - a2 + a3
            const = a1 * c1 + a2 * b2 + a3 * b3

            if coeff_x != 0:
                x = -const / coeff_x
                loesung = f"x = {fmt(x)}"
                erklaerung = (
                    f"Ausmultiplizieren und zusammenfassen: {coeff_x}x + {const} = 0"
                )
            else:
                loesung = "Keine eindeutige Lösung"
                erklaerung = "Die x-Terme heben sich auf"

        else:
            # Schwere Gleichung: Mit Brüchen und Dezimalzahlen
            a1, b1 = random.randint(2, 4), random.randint(1, 3)
            dec = random.choice([1.5, 2.5, 0.5])
            frac_num, frac_den = random.randint(1, 3), random.randint(2, 4)
            c1, d1, e1 = (
                random.randint(2, 8),
                random.randint(5, 20),
                random.randint(2, 8),
            )

            aufgabe = (
                f"{a1}({b1}x - {dec}) + {frac_num}/{frac_den}·(x + {c1}) = {d1}/{e1}"
            )

            # Konvertiere alles zu Brüchen für genaue Berechnung
            dec_frac = Fraction(str(dec))
            frac = Fraction(frac_num, frac_den)
            right = Fraction(d1, e1)

            # a1*b1*x - a1*dec + frac*x + frac*c1 = right
            coeff_x = Fraction(a1 * b1) + frac
            const = -Fraction(a1) * dec_frac + frac * Fraction(c1)

            x = (right - const) / coeff_x
            loesung = f"x = {fmt(float(x), 3)}"
            erklaerung = f"Mit Brüchen auflösen: {fmt(float(coeff_x), 3)}x + {fmt(float(const), 3)} = {fmt(float(right), 3)}"

        return aufgabe, loesung, erklaerung

    def generate_textaufgabe(self, punkte: int) -> tuple[str, str, str]:
        """Generiert realistische Textaufgabe."""
        if punkte <= 10:
            return self._textaufgabe_mittel()
        else:
            return self._textaufgabe_schwer()

    def _textaufgabe_mittel(self) -> tuple[str, str, str]:
        """Mittlere Textaufgabe aus Handwerk/Technik."""
        templates = [
            self._template_gehalt,
            self._template_material,
            self._template_produktion,
            self._template_energie,
        ]
        return random.choice(templates)()

    def _template_gehalt(self) -> tuple[str, str, str]:
        """Gehaltsberechnung."""
        while True:
            beruf = random.choice(list(self.austrian_data.berufe.keys()))
            beruf_data = self.austrian_data.berufe[beruf]
            gehalt = random.randint(beruf_data["gehalt_min"], beruf_data["gehalt_max"])
            stunden_alt = random.randint(*beruf_data["stunden"])
            stunden_neu = random.randint(30, stunden_alt - 2)
            if self._register_task("text/gehalt", [gehalt, stunden_alt, stunden_neu]):
                break

        aufgabe = (
            f"Ein {beruf.replace('_', '-').title()} verdient {gehalt}€ bei {stunden_alt} Stunden/Woche. "
            f"Bei einer Reduktion auf {stunden_neu} Stunden/Woche (gleicher Stundenlohn): "
            f"a) Wie hoch ist das neue Gehalt? b) Um wie viel Prozent sinkt das Gehalt?"
        )

        stundenlohn = gehalt / stunden_alt
        neues_gehalt = stundenlohn * stunden_neu
        prozent = ((gehalt - neues_gehalt) / gehalt) * 100

        loesung = f"a) {fmt(neues_gehalt)}€, b) -{fmt(prozent, 1)}%"
        erklaerung = (
            f"Stundenlohn = {gehalt}/{stunden_alt} = {fmt(stundenlohn)}€/h. "
            f"Neues Gehalt = {fmt(stundenlohn)}×{stunden_neu} = {fmt(neues_gehalt)}€. "
            f"Prozentuale Änderung = ({gehalt} - {fmt(neues_gehalt)})/{gehalt}×100 = {fmt(prozent, 1)}%."
        )

        return aufgabe, loesung, erklaerung

    def _template_material(self) -> tuple[str, str, str]:
        """Materialverbrauch."""
        while True:
            material = random.choice(list(self.austrian_data.materialien.keys()))
            material_data = self.austrian_data.materialien[material]
            laenge = random.randint(200, 500)
            breite = random.randint(10, 30)
            hoehe = random.randint(5, 20)
            preis_map = {
                "stahl": "stahl_kg",
                "aluminium": "aluminium_kg",
                "kupfer": "kupfer_kg",
                "beton": "beton_m3",
                "holz_fichte": None,
            }
            mat_key = material.split("_")[0]
            preis_key = preis_map.get(mat_key)
            if preis_key is None:
                material = "stahl"
                material_data = self.austrian_data.materialien[material]
                preis_key = "stahl_kg"
            if self._register_task(
                "text/material",
                [laenge, breite, hoehe, self.austrian_data.preise[preis_key]],
            ):
                break

        preis_wert = self.austrian_data.preise[preis_key]
        einheit = "€/kg" if preis_key.endswith("_kg") else "€/m³"

        aufgabe = (
            f"Ein Werkstück aus {material.replace('_', ' ').title()} hat die Maße "
            f"{laenge}cm × {breite}cm × {hoehe}cm. "
            f"Dichte: {material_data['dichte']} {material_data['einheit']}. "
            f"Preis: {preis_wert}{einheit}. Berechnen Sie: a) Gewicht b) Materialkosten"
        )

        volumen_cm3 = laenge * breite * hoehe
        volumen_dm3 = volumen_cm3 / 1000
        volumen_m3 = volumen_cm3 / 1_000_000
        gewicht = volumen_dm3 * material_data["dichte"]
        if preis_key.endswith("_kg"):
            kosten = gewicht * preis_wert
        else:
            kosten = volumen_m3 * preis_wert

        loesung = f"a) {fmt(gewicht)}kg, b) {fmt(kosten)}€"
        erklaerung = (
            f"Volumen = {laenge}×{breite}×{hoehe} = {volumen_cm3}cm³ = {fmt(volumen_dm3)}dm³. "
            f"Gewicht = {fmt(volumen_dm3)}×{material_data['dichte']} = {fmt(gewicht)}kg. "
            f"Kosten = {fmt(gewicht)}×{preis_wert} = {fmt(kosten)}€"
        )

        return aufgabe, loesung, erklaerung

    def _template_produktion(self) -> tuple[str, str, str]:
        """Produktionsaufgabe."""
        while True:
            maschinen = random.randint(3, 8)
            zeit = random.randint(20, 50)
            neue_maschinen = random.randint(maschinen + 2, maschinen * 2)
            if self._register_task(
                "text/produktion", [maschinen, zeit, neue_maschinen]
            ):
                break

        aufgabe = (
            f"{maschinen} CNC-Fräsmaschinen produzieren einen Auftrag in {zeit} Stunden. "
            f"Wie lange brauchen {neue_maschinen} Maschinen für denselben Auftrag? "
            f"(Geben Sie das Ergebnis in Stunden und Minuten an)"
        )

        neue_zeit = (maschinen * zeit) / neue_maschinen
        stunden = int(neue_zeit)
        minuten = int((neue_zeit - stunden) * 60)

        loesung = f"{stunden}h {minuten}min"
        erklaerung = (
            f"Auftragsmenge bleibt gleich: {maschinen}×{zeit} = {maschinen * zeit} Maschinenstunden. "
            f"Neue Zeit = {maschinen * zeit}/{neue_maschinen} = {fmt(neue_zeit)}h"
        )

        return aufgabe, loesung, erklaerung

    def _template_energie(self) -> tuple[str, str, str]:
        """Energieverbrauch einer Maschine."""
        while True:
            leistung = round(random.uniform(1.5, 5.0), 1)
            stunden = random.randint(4, 8)
            if self._register_task("text/energie", [int(leistung * 10), stunden]):
                break

        preis = self.austrian_data.preise["strom_kwh"]
        aufgabe = (
            f"Eine Maschine mit {leistung}kW läuft {stunden} Stunden. "
            f"Strompreis: {preis}€/kWh. a) Energieverbrauch? b) Kosten?"
        )

        verbrauch = leistung * stunden
        kosten = verbrauch * preis

        loesung = f"a) {fmt(verbrauch)}kWh, b) {fmt(kosten)}€"
        erklaerung = (
            f"Verbrauch = {leistung}×{stunden} = {fmt(verbrauch)}kWh. "
            f"Kosten = {fmt(verbrauch)}×{preis} = {fmt(kosten)}€"
        )

        return aufgabe, loesung, erklaerung

    def _textaufgabe_schwer(self) -> tuple[str, str, str]:
        """Schwere mehrstufige Textaufgabe."""
        templates = [
            self._template_pumpsystem,
            self._template_mischung,
            self._template_logistik,
            self._template_personalplanung,
        ]
        return random.choice(templates)()

    def _template_pumpsystem(self) -> tuple[str, str, str]:
        """Komplexes Pumpsystem."""
        while True:
            tank = random.randint(1000, 3000)
            pumpe_a_zeit = random.randint(20, 40)
            pumpe_b_zeit = random.randint(30, 60)
            fuellstand = random.randint(40, 70)
            if self._register_task(
                "text/pumpsystem", [tank, pumpe_a_zeit, pumpe_b_zeit, fuellstand]
            ):
                break

        aufgabe = (
            f"Ein Tank fasst {tank}L. Pumpe A füllt ihn in {pumpe_a_zeit}min, "
            f"Pumpe B in {pumpe_b_zeit}min. "
            f"a) Wie lange brauchen beide zusammen? "
            f"b) Bei {fuellstand}% Füllung - wie lange noch mit beiden Pumpen? "
            f"c) Reicht der Tank für 5 Maschinen à 25L/h für 8h Betrieb?"
        )

        rate_a = tank / pumpe_a_zeit
        rate_b = tank / pumpe_b_zeit
        rate_gesamt = rate_a + rate_b

        zeit_gesamt = tank / rate_gesamt
        restmenge = tank * (1 - fuellstand / 100)
        restzeit = restmenge / rate_gesamt

        verbrauch = 5 * 25 * 8
        reicht = "Ja" if tank >= verbrauch else "Nein"

        loesung = f"a) {fmt(zeit_gesamt, 1)}min, b) {fmt(restzeit, 1)}min, c) {reicht} ({verbrauch}L Bedarf)"
        erklaerung = (
            f"Rate A = {tank}/{pumpe_a_zeit} = {fmt(rate_a, 1)}L/min, "
            f"Rate B = {tank}/{pumpe_b_zeit} = {fmt(rate_b, 1)}L/min, "
            f"gemeinsam {fmt(rate_gesamt, 1)}L/min. "
            f"Restmenge = {tank}×(1-{fuellstand}/100) = {fmt(restmenge, 1)}L"
        )

        return aufgabe, loesung, erklaerung

    def _template_mischung(self) -> tuple[str, str, str]:
        """Mischungsaufgabe."""
        while True:
            sorte_a_preis = random.randint(80, 120) / 100
            sorte_b_preis = random.randint(150, 200) / 100
            menge_a = random.randint(20, 40)
            menge_b = random.randint(10, 30)
            zielpreis = (sorte_a_preis + sorte_b_preis) / 2
            x = menge_a * (zielpreis - sorte_a_preis) / (sorte_b_preis - zielpreis)
            if x >= 0 and self._register_task(
                "text/mischung", [menge_a, menge_b, int(x * 10)]
            ):
                break

        aufgabe = (
            f"Eine Metalllegierung wird aus zwei Sorten gemischt: "
            f"Sorte A kostet {sorte_a_preis}€/kg, Sorte B kostet {sorte_b_preis}€/kg. "
            f"Für eine Mischung werden {menge_a}kg von A und {menge_b}kg von B verwendet. "
            f"a) Gesamtkosten? b) Durchschnittspreis pro kg? "
            f"c) Wie viel kg von B für einen Durchschnittspreis von {fmt(zielpreis)}€/kg?"
        )

        kosten_a = sorte_a_preis * menge_a
        kosten_b = sorte_b_preis * menge_b
        gesamtkosten = kosten_a + kosten_b
        gesamtmenge = menge_a + menge_b
        durchschnitt = gesamtkosten / gesamtmenge

        loesung = (
            f"a) {fmt(gesamtkosten)}€, b) {fmt(durchschnitt)}€/kg, c) {fmt(x, 1)}kg"
        )
        erklaerung = (
            f"Gesamtkosten = {fmt(kosten_a)} + {fmt(kosten_b)} = {fmt(gesamtkosten)}€. "
            f"Durchschnitt = {fmt(gesamtkosten)}/{gesamtmenge} = {fmt(durchschnitt)}€/kg"
        )

        return aufgabe, loesung, erklaerung

    def _template_logistik(self) -> tuple[str, str, str]:
        """Logistikaufgabe."""
        while True:
            lkw_kapazitaet = random.randint(8000, 12000)
            paletten = random.randint(20, 30)
            gewicht_palette = random.randint(200, 400)
            strecke = random.randint(200, 500)
            verbrauch = random.randint(25, 35)
            if self._register_task(
                "text/logistik",
                [lkw_kapazitaet, paletten, gewicht_palette, strecke, verbrauch],
            ):
                break

        diesel_preis = self.austrian_data.preise["diesel_l"]
        aufgabe = (
            f"Ein LKW (Nutzlast {lkw_kapazitaet}kg) soll {paletten} Paletten à {gewicht_palette}kg "
            f"über {strecke}km transportieren. Verbrauch: {verbrauch}L/100km, Diesel: {diesel_preis}€/L. "
            f"a) Wie viele Fahrten? b) Gesamte Dieselkosten? c) Kosten pro Palette?"
        )

        gesamtgewicht = paletten * gewicht_palette
        fahrten = math.ceil(gesamtgewicht / lkw_kapazitaet)

        diesel_gesamt = fahrten * strecke * verbrauch / 100
        diesel_kosten = diesel_gesamt * diesel_preis
        kosten_pro_palette = diesel_kosten / paletten

        loesung = f"a) {fahrten} Fahrten, b) {fmt(diesel_kosten)}€, c) {fmt(kosten_pro_palette)}€/Palette"
        erklaerung = (
            f"Gesamtgewicht = {paletten}×{gewicht_palette} = {fmt(gesamtgewicht, 0, True)}kg. "
            f"Diesel = {fahrten}×{strecke}×{verbrauch}/100 = {fmt(diesel_gesamt, 1)}L"
        )

        return aufgabe, loesung, erklaerung

    def _template_personalplanung(self) -> tuple[str, str, str]:
        """Personalplanung für ein Projekt."""
        while True:
            arbeitsstunden = random.randint(400, 800)
            ziel_tage = random.randint(10, 20)
            stunden_tag = random.randint(6, 8)
            arbeiter_vorhanden = random.randint(5, 10)
            ausfall = random.randint(1, 3)
            if self._register_task(
                "text/personal",
                [arbeitsstunden, ziel_tage, stunden_tag, arbeiter_vorhanden, ausfall],
            ):
                break

        aufgabe = (
            f"Ein Projekt umfasst {arbeitsstunden} Arbeitsstunden und soll in {ziel_tage} Tagen "
            f"mit Arbeitstagen zu {stunden_tag} Stunden fertig sein. "
            f"a) Wie viele Arbeiter sind nötig? b) Stehen nur {arbeiter_vorhanden} Arbeiter zur Verfügung, "
            f"wie viele Tage dauert es? c) Fallen {ausfall} Arbeiter aus, wie lange dauert es dann?"
        )

        arbeiter_noetig = math.ceil(arbeitsstunden / (ziel_tage * stunden_tag))
        dauer_b = arbeitsstunden / (arbeiter_vorhanden * stunden_tag)
        dauer_c = arbeitsstunden / ((arbeiter_vorhanden - ausfall) * stunden_tag)

        loesung = f"a) {arbeiter_noetig} Arbeiter, b) {fmt(dauer_b, 1)} Tage, c) {fmt(dauer_c, 1)} Tage"
        erklaerung = (
            f"Tagesleistung pro Arbeiter = {stunden_tag}h. "
            f"Benötigte Arbeiter = {arbeitsstunden}/({ziel_tage}×{stunden_tag}) = {arbeiter_noetig}. "
            f"Mit {arbeiter_vorhanden} Arbeitern: {arbeitsstunden}/({arbeiter_vorhanden}×{stunden_tag}) = {fmt(dauer_b, 1)} Tage. "
            f"Nach Ausfall: {arbeitsstunden}/(({arbeiter_vorhanden - ausfall})×{stunden_tag}) = {fmt(dauer_c, 1)} Tage"
        )

        return aufgabe, loesung, erklaerung

    def generate_stellenwerttabelle(self) -> tuple[str, str, str]:
        """Generiert Stellenwerttabellen-Aufgabe."""
        werte = []
        loesungen = []

        # 1. Ausgeschriebene Zahl
        tausender = random.randint(1, 9)
        hunderter = random.randint(0, 9)
        zehner = random.randint(0, 9)
        einer = random.randint(0, 9)

        zahl = tausender * 1000 + hunderter * 100 + zehner * 10 + einer
        zahl_text = self._zahl_zu_text(zahl)
        werte.append(zahl_text)
        loesungen.append(f"{tausender}T {hunderter}H {zehner}Z {einer}E")

        # 2. Dezimalzahl
        dezimal = round(random.uniform(0.001, 9.999), 3)
        werte.append(fmt(dezimal, 3))
        # Zerlegung in Stellenwerte
        ganz = int(dezimal)
        rest = dezimal - ganz
        z = int(rest * 10)
        h = int(rest * 100) % 10
        t = int(rest * 1000) % 10
        loesungen.append(f"{ganz}E {z}z {h}h {t}t")

        # 3. Bruch
        zaehler = random.randint(100, 999)
        nenner = random.choice([10, 100, 1000])
        werte.append(f"{zaehler}/{nenner}")
        dezimalwert = zaehler / nenner
        loesungen.append(
            f"{fmt(dezimalwert, 3)} → {int(dezimalwert)}E + Dezimalstellen"
        )

        # 4. Gemischte Darstellung
        ganz = random.randint(1, 99)
        dez = random.randint(1, 99)
        werte.append(f"{ganz} und {dez} Hundertstel")
        loesungen.append(f"{ganz}E {dez // 10}z {dez % 10}h")

        aufgabe = "Tragen Sie in die Stellenwerttabelle ein:\n"
        for i, wert in enumerate(werte, 1):
            aufgabe += f"{i}. {wert}\n"

        # Lösung formatieren
        loesung = "Stellenwerttabelle:\n"
        loesung += "| Nr | HT | ZT | T | H | Z | E | , | z | h | t |\n"
        loesung += "|----|----|----|----|----|----|----|----|----|----|----|\n"
        for i, wert in enumerate(loesungen, 1):
            loesung += f"| {i}. | {wert} |\n"

        erklaerung = "HT=Hunderttausender, ZT=Zehntausender, T=Tausender, H=Hunderter, Z=Zehner, E=Einer, z=Zehntel, h=Hundertstel, t=Tausendstel"

        return aufgabe, loesung, erklaerung

    def _zahl_zu_text(self, zahl: int) -> str:
        """Wandelt Zahl in ausgeschriebenen Text."""
        einer_namen = [
            "",
            "ein",
            "zwei",
            "drei",
            "vier",
            "fünf",
            "sechs",
            "sieben",
            "acht",
            "neun",
        ]
        zehner_namen = [
            "",
            "zehn",
            "zwanzig",
            "dreißig",
            "vierzig",
            "fünfzig",
            "sechzig",
            "siebzig",
            "achtzig",
            "neunzig",
        ]
        spezial = {
            11: "elf",
            12: "zwölf",
            13: "dreizehn",
            14: "vierzehn",
            15: "fünfzehn",
            16: "sechzehn",
            17: "siebzehn",
            18: "achtzehn",
            19: "neunzehn",
        }

        if zahl == 0:
            return "null"
        elif zahl < 10:
            return einer_namen[zahl] if zahl != 1 else "eins"
        elif 11 <= zahl <= 19:
            return spezial[zahl]
        elif zahl < 100:
            z = zahl // 10
            e = zahl % 10
            if e == 0:
                return zehner_namen[z]
            elif e == 1:
                return "einund" + zehner_namen[z]
            else:
                return einer_namen[e] + "und" + zehner_namen[z]
        elif zahl < 1000:
            h = zahl // 100
            rest = zahl % 100
            result = einer_namen[h] + "hundert"
            if rest > 0:
                result += self._zahl_zu_text(rest)
            return result
        elif zahl < 10000:
            t = zahl // 1000
            rest = zahl % 1000
            if t == 1:
                result = "eintausend"
            else:
                result = einer_namen[t] + "tausend"
            if rest > 0:
                result += self._zahl_zu_text(rest)
            return result
        else:
            # Für größere Zahlen
            return str(zahl)

    def generate_drei_ansichten(self) -> tuple[str, str, str]:
        """Generiert Drei-Ansichten-Aufgabe."""
        koerper_ascii = {
            "Quader mit Aussparung": (
                "```\n" "┌───────┐\n" "│ ┌───┐ │\n" "│ └───┘ │\n" "└───────┘\n" "```"
            ),
            "L-förmiger Körper": (
                "```\n"
                "    ┌───┐\n"
                "    │   │\n"
                "┌───┼───┘\n"
                "│   │\n"
                "└───┘\n"
                "```"
            ),
            "T-förmiger Körper": (
                "```\n"
                "┌───┬───┬───┐\n"
                "│   │   │   │\n"
                "└───┼───┼───┘\n"
                "    │   │\n"
                "    └───┘\n"
                "```"
            ),
            "Treppenförmiger Körper": (
                "```\n" "┌───┐\n" "│   └───┐\n" "│       │\n" "└───────┘\n" "```"
            ),
            "Würfel mit quadratischer Bohrung": (
                "```\n" "┌───┐\n" "│┌─┐│\n" "│└─┘│\n" "└───┘\n" "```"
            ),
            "U-förmiger Körper": (
                "```\n"
                "┌───┐ ┌───┐\n"
                "│   │ │   │\n"
                "└───┼─┼───┘\n"
                "    │ │\n"
                "    └─┘\n"
                "```"
            ),
            "Z-förmiger Körper": (
                "```\n" "┌───┐\n" "└─┐ │\n" "  │ └─┐\n" "  └───┘\n" "```"
            ),
        }

        koerper = random.choice(list(koerper_ascii.keys()))

        aufgabe = (
            f"Skizzieren Sie den {koerper} in Vorderansicht, Seitenansicht (von links) und Draufsicht.\n"
            "Ordnen Sie die Ansichten nach technischer Norm an.\n"
            "(Verwenden Sie einen weichen Bleistift, Lineal ist nicht notwendig)"
        )

        aufgabe += "\n" + koerper_ascii.get(koerper, "")

        loesung = f"Drei Ansichten des {koerper} nach DIN/ISO"
        erklaerung = (
            "Anordnung nach 1. Winkelprojektion (DIN/ISO): Draufsicht über der Vorderansicht, "
            "Seitenansicht (linke Ansicht) links der Vorderansicht; gleiche Maßstäbe beachten."
        )

        return aufgabe, loesung, erklaerung

    def generate_koerpernetz(self) -> tuple[str, str, str]:
        """Generiert Körpernetz-Aufgabe."""
        koerper_typen = {
            "Würfel": 6,
            "Quader": 6,
            "Pyramide (quadratische Grundfläche)": 5,
            "Prisma (dreieckige Grundfläche)": 5,
            "Tetraeder": 4,
        }

        netz_ascii = {
            "Würfel": (
                "```\n"
                "    ┌───┐\n"
                "┌───┼───┼───┐\n"
                "└───┼───┼───┘\n"
                "    └───┘\n"
                "```"
            ),
            "Quader": (
                "```\n"
                "    ┌─────┐\n"
                "┌─────┼─────┼─────┐\n"
                "└─────┼─────┼─────┘\n"
                "    └─────┘\n"
                "```"
            ),
            "Pyramide (quadratische Grundfläche)": (
                "```\n" "  ▲\n" " ▲▲▲\n" "▲▲▲▲▲\n" "  ◼\n" "```"
            ),
        }

        koerper = random.choice(list(koerper_typen.keys()))
        flaechen = koerper_typen[koerper]

        aufgabe = f"Skizzieren Sie das Körpernetz eines {koerper}.\n"
        aufgabe += f"Beachten Sie: Der Körper hat {flaechen} Flächen."

        if koerper in netz_ascii:
            aufgabe += "\n" + netz_ascii[koerper]

        loesung = f"Körpernetz des {koerper} mit {flaechen} Flächen"
        erklaerung = "Alle Flächen müssen zusammenhängend und ausklappbar sein"

        return aufgabe, loesung, erklaerung

    def generate_runden(self) -> tuple[str, str, str]:
        """Generiert Rundungsaufgabe."""
        zahlen = []
        stellen = []
        gerundete = []

        for _ in range(4):
            # Verschiedene Zahlentypen
            if random.random() < 0.5:
                # Dezimalzahl
                zahl = round(random.uniform(0.001, 9999.999), 4)
                stelle = random.choice(["E", "z", "h", "t"])
            else:
                # Große Zahl
                zahl = random.randint(10000, 999999) + random.random()
                stelle = random.choice(["Z", "H", "T", "ZT", "HT"])

            zahlen.append(zahl)
            stellen.append(stelle)

            # Runden
            gerundet_wert = self.math_solver.round_to_place(zahl, stelle)
            gerundete.append(gerundet_wert)

        aufgabe = "Runden Sie auf die angegebene Stelle:\n"
        loesung = ""

        for i, (zahl, stelle, gerundet) in enumerate(
            zip(zahlen, stellen, gerundete), 1
        ):
            # Formatierung der Zahl je nach Größe
            if zahl >= 1000:
                zahl_str = fmt(zahl, 2)
            else:
                zahl_str = fmt(zahl, 4)

            aufgabe += f"{i}. {zahl_str} (≈{stelle}) = _____\n"

            if isinstance(gerundet, int) or float(gerundet).is_integer():
                loesung += f"{i}. {fmt(gerundet)}\n"
            elif gerundet >= 1000:
                loesung += f"{i}. {fmt(gerundet, 0)}\n"
            else:
                loesung += f"{i}. {fmt(gerundet, 4)}".rstrip("0").rstrip(",") + "\n"

        erklaerung = "E=Einer, z=Zehntel, h=Hundertstel, t=Tausendstel, Z=Zehner, H=Hunderter, T=Tausender, ZT=Zehntausender, HT=Hunderttausender"

        return aufgabe, loesung, erklaerung

    def generate_einheiten(self, niveau: int) -> tuple[str, str, str]:
        """Generiert Einheitenumwandlung."""
        if niveau == 1:  # Leicht
            conversions = [
                (random.randint(3, 20), "m", "cm"),
                (random.randint(2, 15), "kg", "g"),
                (random.randint(1, 10), "l", "ml"),
            ]
        elif niveau == 2:  # Mittel
            conversions = [
                (random.randint(2, 10), "m²", "cm²"),
                (random.randint(5, 30), "dm³", "l"),
                (random.randint(100, 500), "cm", "dm"),
            ]
        else:  # Schwer
            conversions = [
                (random.randint(500000, 2000000), "cm³", "m³"),
                (random.randint(180, 420), "min", "h"),
                (random.uniform(0.1, 9.9), "km²", "ha"),
            ]

        aufgabe = ""
        loesung = ""

        for i, (wert, von, nach) in enumerate(conversions, 1):
            aufgabe += f"{i}. {wert} {von} = _____ {nach}\n"
            ergebnis = self.converter.convert(wert, von, nach)

            if ergebnis is not None:
                # Formatierung ohne wissenschaftliche Notation
                if ergebnis >= 10000:
                    loesung += f"{i}. {fmt(ergebnis, 0)} {nach}\n"
                elif ergebnis >= 1:
                    loesung += f"{i}. {fmt(ergebnis)} {nach}\n"
                elif ergebnis >= 0.01:
                    loesung += f"{i}. {fmt(ergebnis, 4)} {nach}\n"
                else:
                    loesung += f"{i}. {fmt(ergebnis, 6)} {nach}\n"
            else:
                loesung += f"{i}. [Konvertierung nicht möglich]\n"

        erklaerung = "Verwenden Sie die korrekten Umrechnungsfaktoren"

        return aufgabe, loesung, erklaerung

    def generate_geometrie(self, zeichnen: bool = True) -> tuple[str, str, str]:
        """Generiert Geometrieaufgabe."""
        if zeichnen:
            # L-förmiges Werkstück
            l1 = random.randint(40, 80)
            w1 = random.randint(30, 60)
            l2 = random.randint(20, 40)
            w2 = random.randint(15, 35)
            massstab = random.choice([10, 20, 50])

            aufgabe = (
                f"Ein L-förmiges Werkstück besteht aus zwei Rechtecken:\n"
                f"Hauptteil: {l1}mm × {w1}mm\n"
                f"Ansatz: {l2}mm × {w2}mm (an der Ecke)\n"
                f"a) Berechnen Sie Umfang und Fläche\n"
                f"b) Zeichnen Sie es im Maßstab 1:{massstab}"
            )

            flaeche = self.geometry.l_shape_area(l1, w1, l2, w2)
            umfang = self.geometry.l_shape_perimeter(l1, w1, l2, w2)

            loesung = f"Fläche: {flaeche}mm², Umfang: {umfang}mm"
            erklaerung = f"Fläche = {l1}×{w1} + {l2}×{w2} = {flaeche}mm²"

        else:
            # Volumenberechnung
            laenge = random.randint(300, 600)
            breite = random.randint(15, 35)
            hoehe = random.randint(8, 20)
            material = random.choice(["stahl", "aluminium"])
            dichte = self.austrian_data.materialien[material]["dichte"]
            preis = self.austrian_data.preise[f"{material}_kg"]
            verschnitt = random.randint(8, 15)

            aufgabe = (
                f"Ein {material.title()}träger: {laenge}cm × {breite}cm × {hoehe}cm\n"
                f"Dichte: {dichte}kg/dm³, Preis: {preis}€/kg\n"
                f"a) Gewicht? b) Materialkosten? c) Mit {verschnitt}% Verschnitt?"
            )

            volumen = self.geometry.cuboid_volume(
                laenge / 10, breite / 10, hoehe / 10
            )  # in dm³
            gewicht = volumen * dichte
            kosten = gewicht * preis
            kosten_verschnitt = kosten * (1 + verschnitt / 100)

            loesung = (
                f"a) {fmt(gewicht)}kg, b) {fmt(kosten)}€, c) {fmt(kosten_verschnitt)}€"
            )
            erklaerung = f"Volumen: {fmt(volumen)}dm³, Gewicht: {fmt(gewicht)}kg"

        return aufgabe, loesung, erklaerung


class TestGenerator:
    """Hauptklasse für Testgenerierung."""

    def __init__(
        self,
        schwierigkeit: Schwierigkeit = Schwierigkeit.MITTEL,
        seed: int | None = None,
    ):
        if seed is not None:
            random.seed(seed)
        self.schwierigkeit = schwierigkeit
        self.generator = AufgabenGenerator(schwierigkeit)
        self.test_content = ""
        self.solutions = ""
        self.detailed_solutions = []

    def generate_complete_test(self) -> tuple[str, str, list[dict]]:
        """Generiert kompletten Test mit exakt 100 Punkten."""

        # Header
        self.test_content = """# Überstiegstest - Technische Basisausbildung

**Name: ________________________________    Datum: ________________**

**Bearbeitungszeit: 90 Minuten**
**Gesamtpunktzahl: 100 Punkte**
**Bestehensgrenze: 60 Punkte**

---

"""

        self.solutions = """# LÖSUNGEN - Überstiegstest

**Lösungsschlüssel für Lehrkraft**

---

"""

        # 1. GRUNDRECHENARTEN (20 Punkte)
        self._add_grundrechenarten()

        # 2. ZAHLENRAUM (20 Punkte)
        self._add_zahlenraum()

        # 3. TEXTAUFGABEN (20 Punkte)
        self._add_textaufgaben()

        # 4. BRÜCHE UND GLEICHUNGEN (20 Punkte)
        self._add_brueche_gleichungen()

        # 5. RAUMVORSTELLUNG (20 Punkte)
        self._add_raumvorstellung()

        # Bewertungsschlüssel
        self._add_bewertung()

        return self.test_content, self.solutions, self.detailed_solutions

    def _add_grundrechenarten(self):
        """Fügt Grundrechenarten-Sektion hinzu."""
        self.test_content += "## 1. Grundrechenarten (20 Punkte)\n"
        self.test_content += "*Beachten Sie: Klammer vor Punkt vor Strich! Runden Sie auf 2 Dezimalstellen.*\n\n"

        self.solutions += "## 1. Grundrechenarten\n\n"

        # 2 leichte (je 2 Punkte)
        for i, punkte in enumerate([2, 2], start=1):
            aufgabe, loesung, erklaerung = self.generator.generate_grundrechnung(punkte)
            self.test_content += f"**a.{i})** {aufgabe} = _____ **(2 Punkte)**\n\n"
            self.solutions += f"**a.{i})** {aufgabe} = **{loesung}**\n"
            self.solutions += f"   {erklaerung}\n\n"
            self.detailed_solutions.append(
                {
                    "nummer": f"1.a.{i}",
                    "aufgabe": aufgabe,
                    "loesung": loesung,
                    "erklaerung": erklaerung,
                    "punkte": punkte,
                }
            )

        # 2 mittlere (je 3 Punkte)
        for i, punkte in enumerate([3, 3], start=1):
            aufgabe, loesung, erklaerung = self.generator.generate_grundrechnung(punkte)
            self.test_content += f"**b.{i})** {aufgabe} = _____ **(3 Punkte)**\n\n"
            self.solutions += f"**b.{i})** {aufgabe} = **{loesung}**\n"
            self.solutions += f"   {erklaerung}\n\n"
            self.detailed_solutions.append(
                {
                    "nummer": f"1.b.{i}",
                    "aufgabe": aufgabe,
                    "loesung": loesung,
                    "erklaerung": erklaerung,
                    "punkte": punkte,
                }
            )

        # 2 schwere (je 5 Punkte)
        for i, punkte in enumerate([5, 5], start=1):
            aufgabe, loesung, erklaerung = self.generator.generate_grundrechnung(punkte)
            self.test_content += f"**c.{i})** {aufgabe} = _____ **(5 Punkte)**\n\n"
            self.solutions += f"**c.{i})** {aufgabe} = **{loesung}**\n"
            self.solutions += f"   {erklaerung}\n\n"
            self.detailed_solutions.append(
                {
                    "nummer": f"1.c.{i}",
                    "aufgabe": aufgabe,
                    "loesung": loesung,
                    "erklaerung": erklaerung,
                    "punkte": punkte,
                }
            )

        self.test_content += "\n---\n\n"

    def _add_zahlenraum(self):
        """Fügt Zahlenraum-Sektion hinzu."""
        self.test_content += "## 2. Zahlenraum (20 Punkte)\n\n"
        self.solutions += "## 2. Zahlenraum\n\n"

        # Zahlenstrahl (5 Punkte)
        self.test_content += "**a) Zahlenstrahl (5 Punkte)**\n"
        self.test_content += (
            "Tragen Sie folgende Werte ein: 0,5; -2,8; 6; 1/2; -3/4\n\n"
        )
        self.test_content += "```\n"
        self.test_content += "-10 _____|_____|_____|_____|_____|_____|_____|_____|_____|_____|_____ +10\n"
        self.test_content += "```\n\n"
        self.solutions += "**a)** Zahlenstrahl mit eingetragenen Werten\n\n"

        # Stellenwerttabelle (5 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_stellenwerttabelle()
        self.test_content += f"**b) Stellenwerttabelle (5 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**b)** {loesung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "2.b",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 5,
            }
        )

        # Runden (3 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_runden()
        self.test_content += f"**c) Runden (3 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**c)** {loesung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "2.c",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 3,
            }
        )

        # Einheiten (7 Punkte: 2+2+3)
        self.test_content += "**d) Einheitenumwandlungen (7 Punkte)**\n\n"

        for niveau, punkte, titel in [
            (1, 2, "Leicht"),
            (2, 2, "Mittel"),
            (3, 3, "Schwer"),
        ]:
            aufgabe, loesung, erklaerung = self.generator.generate_einheiten(niveau)
            self.test_content += f"*{titel} ({punkte} Punkte):*\n{aufgabe}\n"
            self.solutions += f"**d.{niveau})** {loesung}\n"
            self.detailed_solutions.append(
                {
                    "nummer": f"2.d.{niveau}",
                    "aufgabe": aufgabe,
                    "loesung": loesung,
                    "erklaerung": erklaerung,
                    "punkte": punkte,
                }
            )

        self.test_content += "\n---\n\n"

    def _add_textaufgaben(self):
        """Fügt Textaufgaben-Sektion hinzu."""
        self.test_content += "## 3. Textaufgaben (20 Punkte)\n\n"
        self.solutions += "## 3. Textaufgaben\n\n"

        # Mittlere Aufgabe (10 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_textaufgabe(10)
        self.test_content += f"**a) (10 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**a)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "3.a",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 10,
            }
        )

        # Schwere Aufgabe (10 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_textaufgabe(15)
        self.test_content += f"**b) (10 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**b)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "3.b",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 10,
            }
        )

        self.test_content += "\n---\n\n"

    def _add_brueche_gleichungen(self):
        """Fügt Brüche und Gleichungen hinzu."""
        self.test_content += "## 4. Brüche und Gleichungen (20 Punkte)\n\n"
        self.solutions += "## 4. Brüche und Gleichungen\n\n"

        # 3 Bruchrechnungen (je 4 Punkte = 12 Punkte)
        self.test_content += (
            "**Bruchrechnung (12 Punkte)**\n*Kürzen Sie vollständig!*\n\n"
        )

        for i in range(3):
            aufgabe, loesung, erklaerung = self.generator.generate_bruchaufgabe(4)
            self.test_content += f"**a.{i + 1})** {aufgabe} = _____ **(4 Punkte)**\n\n"
            self.solutions += (
                f"**a.{i + 1})** {aufgabe} = **{loesung}**\n   {erklaerung}\n\n"
            )
            self.detailed_solutions.append(
                {
                    "nummer": f"4.a.{i + 1}",
                    "aufgabe": aufgabe,
                    "loesung": loesung,
                    "erklaerung": erklaerung,
                    "punkte": 4,
                }
            )

        # 2 Gleichungen (je 4 Punkte = 8 Punkte)
        self.test_content += "**Gleichungen (8 Punkte)**\n\n"

        # Mittlere Gleichung
        aufgabe, loesung, erklaerung = self.generator.generate_gleichung(schwer=False)
        self.test_content += f"**b.1)** {aufgabe} **(4 Punkte)**\n\n"
        self.solutions += f"**b.1)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "4.b.1",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 4,
            }
        )

        # Schwere Gleichung
        aufgabe, loesung, erklaerung = self.generator.generate_gleichung(schwer=True)
        self.test_content += f"**b.2)** {aufgabe} **(4 Punkte)**\n\n"
        self.solutions += f"**b.2)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "4.b.2",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 4,
            }
        )

        self.test_content += "\n---\n\n"

    def _add_raumvorstellung(self):
        """Fügt Raumvorstellung-Sektion hinzu."""
        self.test_content += "## 5. Raumvorstellung (20 Punkte)\n\n"
        self.solutions += "## 5. Raumvorstellung\n\n"

        # Drei-Ansichten (5 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_drei_ansichten()
        self.test_content += f"**a) Drei Ansichten (5 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**a)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "5.a",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 5,
            }
        )

        # Körpernetz (5 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_koerpernetz()
        self.test_content += f"**b) Körpernetz (5 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**b)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "5.b",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 5,
            }
        )

        # Geometrie Zeichnung (5 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_geometrie(zeichnen=True)
        self.test_content += (
            f"**c) Geometrische Berechnung und Zeichnung (5 Punkte)**\n{aufgabe}\n\n"
        )
        self.solutions += f"**c)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "5.c",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 5,
            }
        )

        # Volumenberechnung (5 Punkte)
        aufgabe, loesung, erklaerung = self.generator.generate_geometrie(zeichnen=False)
        self.test_content += f"**d) Volumen und Gewicht (5 Punkte)**\n{aufgabe}\n\n"
        self.solutions += f"**d)** {loesung}\n   {erklaerung}\n\n"
        self.detailed_solutions.append(
            {
                "nummer": "5.d",
                "aufgabe": aufgabe,
                "loesung": loesung,
                "erklaerung": erklaerung,
                "punkte": 5,
            }
        )

    def _add_bewertung(self):
        """Fügt Bewertungsschlüssel hinzu."""
        bewertung = """
---

## Bewertungsschlüssel

| Note | Bezeichnung | Punkte | Prozent |
|------|------------|--------|---------|
| 1 | Sehr gut | 90-100 | 90-100% |
| 2 | Gut | 80-89 | 80-89% |
| 3 | Befriedigend | 70-79 | 70-79% |
| 4 | Genügend | 60-69 | 60-69% |
| 5 | Nicht genügend | 0-59 | 0-59% |

**Viel Erfolg!**
"""
        self.test_content += bewertung

        self.solutions += "\n## Punkteverteilung\n\n"
        self.solutions += "- Grundrechenarten: 20 Punkte (2+2+3+3+5+5)\n"
        self.solutions += "- Zahlenraum: 20 Punkte (5+5+3+7)\n"
        self.solutions += "- Textaufgaben: 20 Punkte (10+10)\n"
        self.solutions += "- Brüche/Gleichungen: 20 Punkte (12+8)\n"
        self.solutions += "- Raumvorstellung: 20 Punkte (5+5+5+5)\n"
        self.solutions += "\n**Gesamt: 100 Punkte**\n"


class OutputManager:
    """Verwaltet verschiedene Ausgabeformate."""

    @staticmethod
    def save_markdown(content: str, filename: str):
        """Speichert als Markdown-Datei."""
        with open(filename, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"✓ Markdown gespeichert: {filename}")

    @staticmethod
    def save_word(
        test_content: str, solutions: str, test_name: str = "ueberstiegstest"
    ):
        """Erstellt professionelle Word-Dokumente."""
        if not HAS_DOCX:
            print(
                "❌ Word-Export nicht verfügbar. Installieren Sie: pip install python-docx"
            )
            return

        # Test-Dokument
        doc = Document()

        # Titel
        title = doc.add_heading("Überstiegstest", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading(
            "Technische Basisausbildung", 1
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header-Informationen
        doc.add_paragraph()
        info = doc.add_paragraph()
        info.add_run("Name: ").bold = True
        info.add_run("_" * 40 + "    ")
        info.add_run("Datum: ").bold = True
        info.add_run("_" * 20)

        doc.add_paragraph("Bearbeitungszeit: 90 Minuten")
        doc.add_paragraph("Gesamtpunktzahl: 100 Punkte")
        doc.add_paragraph("Bestehensgrenze: 60 Punkte")

        # Seitenumbruch
        doc.add_page_break()

        # Test-Inhalt parsen und formatieren
        OutputManager._parse_markdown_to_word(doc, test_content)

        # Speichern
        doc.save(f"{test_name}.docx")
        print(f"✓ Word-Test gespeichert: {test_name}.docx")

        # Lösungsdokument
        sol_doc = Document()
        sol_title = sol_doc.add_heading("LÖSUNGEN - Überstiegstest", 0)
        sol_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sol_doc.add_paragraph()
        p = sol_doc.add_paragraph()
        r = p.add_run("Lösungsschlüssel für Lehrkraft")
        r.bold = True
        sol_doc.add_paragraph()

        OutputManager._parse_markdown_to_word(sol_doc, solutions)

        sol_doc.save(f"{test_name}_loesungen.docx")
        print(f"✓ Word-Lösungen gespeichert: {test_name}_loesungen.docx")

    @staticmethod
    def _parse_markdown_to_word(doc, markdown_text: str):
        """Konvertiert Markdown zu Word-Format."""
        if not HAS_DOCX:
            return

        lines = markdown_text.split("\n")
        current_table = None

        for line in lines:
            if line.startswith("## "):
                # Hauptüberschrift
                doc.add_heading(line[3:], 2)
            elif line.startswith("### "):
                # Unterüberschrift
                doc.add_heading(line[4:], 3)
            elif line.startswith("**") and line.endswith("**"):
                # Fett
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            elif (
                line.startswith("*")
                and line.endswith("*")
                and not line.startswith("**")
            ):
                # Kursiv
                p = doc.add_paragraph()
                p.add_run(line[1:-1]).italic = True
            elif line.startswith("- "):
                # Aufzählung
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("```"):
                continue
            elif line.startswith("|"):
                cols = [c.strip() for c in line.strip("|").split("|")]
                if all(set(c) <= {"-", " "} for c in cols):
                    continue
                if current_table is None:
                    current_table = doc.add_table(rows=1, cols=len(cols))
                    hdr = current_table.rows[0].cells
                    for j, txt in enumerate(cols):
                        hdr[j].text = txt
                else:
                    row = current_table.add_row().cells
                    for j, txt in enumerate(cols):
                        row[j].text = txt
            elif line.strip() == "---":
                # Horizontale Linie
                doc.add_paragraph("_" * 50)
            elif line.strip():
                current_table = None
                # Normaler Text
                doc.add_paragraph(line)
            else:
                current_table = None

    @staticmethod
    def save_latex(
        test_content: str, solutions: str, test_name: str = "ueberstiegstest"
    ):
        """Erstellt LaTeX-Dokumente."""
        latex_test = OutputManager._markdown_to_latex(test_content, is_solution=False)
        latex_solutions = OutputManager._markdown_to_latex(solutions, is_solution=True)

        # Speichere LaTeX-Dateien
        with open(f"{test_name}.tex", "w", encoding="utf-8") as f:
            f.write(latex_test)
        print(f"✓ LaTeX-Test gespeichert: {test_name}.tex")

        with open(f"{test_name}_loesungen.tex", "w", encoding="utf-8") as f:
            f.write(latex_solutions)
        print(f"✓ LaTeX-Lösungen gespeichert: {test_name}_loesungen.tex")

        # Versuche PDF zu erstellen
        if HAS_LATEX:
            try:
                subprocess.run(["pdflatex", f"{test_name}.tex"], capture_output=True)
                subprocess.run(
                    ["pdflatex", f"{test_name}_loesungen.tex"], capture_output=True
                )
                print(f"✓ PDFs erstellt: {test_name}.pdf, {test_name}_loesungen.pdf")
            except Exception:
                print("⚠ PDF-Erstellung fehlgeschlagen")

    @staticmethod
    def _markdown_to_latex(markdown_text: str, is_solution: bool = False) -> str:
        """Konvertiert Markdown zu LaTeX."""
        latex = r"""\documentclass[12pt,a4paper]{article}
\usepackage[utf8]{inputenc}
\usepackage[ngerman]{babel}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{geometry}
\usepackage{fancyhdr}
\usepackage{graphicx}
\usepackage{tikz}

\geometry{margin=2.5cm}
\pagestyle{fancy}
\fancyhf{}
\rfoot{\thepage}
"""

        if is_solution:
            latex += r"\lhead{LÖSUNGEN - Überstiegstest}"
        else:
            latex += r"\lhead{Überstiegstest - Technische Basisausbildung}"

        latex += r"""
\begin{document}

"""

        if not is_solution:
            latex += r"""\begin{center}
\Large\textbf{Überstiegstest}\\
\large Technische Basisausbildung\\[2em]
\normalsize
Name: \underline{\hspace{8cm}} \quad Datum: \underline{\hspace{4cm}}\\[1em]
Bearbeitungszeit: 90 Minuten\\
Gesamtpunktzahl: 100 Punkte\\
Bestehensgrenze: 60 Punkte
\end{center}

\vspace{2em}

"""
        else:
            latex += r"""\begin{center}
\Large\textbf{LÖSUNGEN - Überstiegstest}\\
\large Lösungsschlüssel für Lehrkraft
\end{center}

\vspace{2em}

"""

        # Konvertiere Markdown zu LaTeX
        lines = markdown_text.split("\n")
        in_list = False
        in_table = False
        in_code = False
        expected_cols = 0

        def _tex_escape(s: str) -> str:
            return (
                s.replace("\\", "\textbackslash{}")
                .replace("&", r"\&")
                .replace("%", r"\%")
                .replace("$", r"\$")
                .replace("#", r"\#")
                .replace("_", r"\_")
                .replace("{", r"\{")
                .replace("}", r"\}")
                .replace("~", "\textasciitilde{}")
                .replace("^", "\textasciicircum{}")
            )

        for line in lines:
            if line.strip().startswith("```"):
                in_code = not in_code
                continue
            if in_code:
                continue
            if line.startswith("## "):
                latex += r"\section{" + _tex_escape(line[3:]) + "}\n"
            elif line.startswith("### "):
                latex += r"\subsection{" + _tex_escape(line[4:]) + "}\n"
            elif line.startswith("**") and line.endswith("**"):
                latex += f"\textbf{{{_tex_escape(line[2:-2])}}}\n\n"
            elif line.startswith("*") and line.endswith("*"):
                latex += f"\textit{{{_tex_escape(line[1:-1])}}}\n\n"
            elif line.startswith("- "):
                if not in_list:
                    latex += r"\begin{itemize}" + "\n"
                    in_list = True
                latex += r"\item " + _tex_escape(line[2:]) + "\n"
                continue
            elif line.startswith("|"):
                cols = [c.strip() for c in line.strip("|").split("|")]
                if all(set(c) <= {"-", " "} for c in cols):
                    continue
                if not in_table:
                    in_table = True
                    expected_cols = len(cols)
                    colspec = " | ".join(["l"] * expected_cols)
                    latex += r"\begin{tabular}{" + colspec + "}\n" + r"\hline" + "\n"
                else:
                    if len(cols) < expected_cols:
                        cols += [""] * (expected_cols - len(cols))
                    elif len(cols) > expected_cols:
                        cols = cols[:expected_cols]
                latex += " & ".join(_tex_escape(c) for c in cols) + " \\\n"
                continue
            else:
                if in_list:
                    latex += r"\end{itemize}" + "\n"
                    in_list = False
                if in_table:
                    latex += r"\hline" + "\n" + r"\end{tabular}" + "\n\n"
                    in_table = False
                if "=" in line and any(
                    op in line for op in ["+", "-", "·", ":", "(", ")"]
                ):
                    equation = line.replace("·", r"\cdot").replace(":", r"\div")
                    latex += f"$${equation}$$\n"
                elif line.strip():
                    latex += _tex_escape(line) + "\n\n"

        if in_list:
            latex += r"\end{itemize}" + "\n"
        if in_table:
            latex += r"\hline" + "\n" + r"\end{tabular}" + "\n"

        latex += r"\end{document}"

        return latex


def main():
    """Hauptprogramm."""
    print(
        """
╔══════════════════════════════════════════════════════════╗
║         ÜBERSTIEGSTEST GENERATOR v2.0                    ║
║         Technische Basisausbildung                       ║
╚══════════════════════════════════════════════════════════╝
"""
    )

    parser = argparse.ArgumentParser()
    parser.add_argument("--stufe", choices=["einfach", "mittel", "schwer"])
    parser.add_argument("--seed", type=int)
    args = parser.parse_args()

    if args.stufe:
        map_stufe = {
            "einfach": Schwierigkeit.EINFACH,
            "mittel": Schwierigkeit.MITTEL,
            "schwer": Schwierigkeit.SCHWER,
        }
        schwierigkeit = map_stufe[args.stufe]
    else:
        if sys.stdin.isatty():
            print("Wählen Sie die Schwierigkeit:")
            print("1. Einfach")
            print("2. Mittel (Standard)")
            print("3. Schwer")
            choice = input("\nIhre Wahl (1-3, Enter für Standard): ").strip()
            if choice == "1":
                schwierigkeit = Schwierigkeit.EINFACH
                print("→ Schwierigkeit: EINFACH")
            elif choice == "3":
                schwierigkeit = Schwierigkeit.SCHWER
                print("→ Schwierigkeit: SCHWER")
            else:
                schwierigkeit = Schwierigkeit.MITTEL
                print("→ Schwierigkeit: MITTEL (Standard)")
        else:
            schwierigkeit = Schwierigkeit.MITTEL
            print("→ Schwierigkeit: MITTEL (Standard)")

    print("\nGeneriere Test...")
    print("-" * 50)

    generator = TestGenerator(schwierigkeit, seed=args.seed)
    test_content, solutions, detailed = generator.generate_complete_test()

    # Ausgabe
    output = OutputManager()

    print("\n" + "=" * 50)
    print("TEST AUSGABE:")
    print("=" * 50)
    print(test_content)

    print("\n" + "=" * 50)
    print("LÖSUNGEN AUSGABE:")
    print("=" * 50)
    print(solutions)

    # Dateien speichern
    print("\n" + "=" * 50)
    print("Speichere Dateien...")
    print("=" * 50)

    # Markdown
    output.save_markdown(test_content, "ueberstiegstest.md")
    output.save_markdown(solutions, "ueberstiegstest_loesungen.md")

    # Word
    output.save_word(test_content, solutions)

    try:
        output.save_latex(test_content, solutions)
    except Exception:
        print("⚠ LaTeX-Export übersprungen")

    # Detaillierte Lösungen als JSON
    with open("ueberstiegstest_details.json", "w", encoding="utf-8") as f:
        json.dump(detailed, f, ensure_ascii=False, indent=2)
    print("✓ Detaillierte Lösungen gespeichert: ueberstiegstest_details.json")

    print("\n" + "=" * 50)
    print("✅ TEST ERFOLGREICH GENERIERT!")
    print("=" * 50)
    print(
        f"""
Erstellte Dateien:
- ueberstiegstest.md (Markdown-Test)
- ueberstiegstest_loesungen.md (Markdown-Lösungen)
- ueberstiegstest.docx (Word-Test)
- ueberstiegstest_loesungen.docx (Word-Lösungen)
- ueberstiegstest_details.json (Detaillierte Lösungen)

Schwierigkeit: {schwierigkeit.name}
Gesamtpunkte: 100
Kategorien: 5 × 20 Punkte
"""
    )


if __name__ == "__main__":
    main()
